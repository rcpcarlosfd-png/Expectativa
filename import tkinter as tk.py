import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import os
import locale
import json

# Configurar locale para español
try:
    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_TIME, 'Spanish_Spain.1252')
    except:
        pass

class DueDiligenceSystem:
    def __init__(self, root):
        self.root = root
        self.root.title("Sistema de Debida Diligencia - Profesional")
        self.root.geometry("1200x750")
        
        # Configurar estilo profesional
        self.configurar_estilos()
        
        # Datos
        self.sujetos = []
        self.plantilla_modificable = self.cargar_plantilla_default()
        self.configuraciones = self.cargar_configuraciones_default()
        self.usuarios_guardados = self.cargar_usuarios_guardados()
        self.serapios_guardados = self.cargar_serapios_guardados()
        self.incisos_personalizados = []
        self.gestion_actual = ""  # Para identificar gestiones
        
        # Cargar sujetos guardados desde archivo
        self.cargar_sujetos_guardados()
        
        # Crear header profesional
        self.crear_header()
        
        # Crear notebook (pestañas)
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill='both', expand=True, padx=15, pady=(0, 15))
        
        # Crear pestañas
        self.crear_tab_datos_generales()
        self.crear_tab_sujetos()
        self.crear_tab_resultados()
        self.crear_tab_plantilla()
        self.crear_tab_configuraciones()
        self.crear_tab_generar()
    
    def configurar_estilos(self):
        """Configura estilos profesionales para la aplicación"""
        style = ttk.Style()
        
        # Tema base
        style.theme_use('clam')
        
        # Colores profesionales
        color_primario = '#2C3E50'
        color_secundario = '#34495E'
        color_acento = '#3498DB'
        color_fondo = '#ECF0F1'
        color_texto = '#2C3E50'
        
        # Configurar Notebook (pestañas)
        style.configure('TNotebook', background=color_fondo, borderwidth=0)
        style.configure('TNotebook.Tab', 
                       background=color_secundario,
                       foreground='white',
                       padding=[20, 10],
                       font=('Segoe UI', 10, 'bold'))
        style.map('TNotebook.Tab',
                 background=[('selected', color_acento)],
                 foreground=[('selected', 'white')])
        
        # Configurar LabelFrame
        style.configure('TLabelframe', 
                       background='white',
                       borderwidth=2,
                       relief='groove')
        style.configure('TLabelframe.Label',
                       background='white',
                       foreground=color_primario,
                       font=('Segoe UI', 11, 'bold'))
        
        # Configurar Botones
        style.configure('TButton',
                       background=color_acento,
                       foreground='white',
                       borderwidth=0,
                       focuscolor='none',
                       padding=[15, 8],
                       font=('Segoe UI', 10))
        style.map('TButton',
                 background=[('active', '#2980B9')])
        
        # Configurar Labels
        style.configure('TLabel',
                       background='white',
                       foreground=color_texto,
                       font=('Segoe UI', 10))
        
        # Configurar Entry
        style.configure('TEntry',
                       fieldbackground='white',
                       borderwidth=1)
        
        # Configurar Combobox
        style.configure('TCombobox',
                       fieldbackground='white',
                       background='white',
                       borderwidth=1)
        
        # Configurar el fondo de la ventana principal
        self.root.configure(bg=color_fondo)
    
    def crear_header(self):
        """Crea un encabezado profesional para la aplicación"""
        header_frame = tk.Frame(self.root, bg='#2C3E50', height=80)
        header_frame.pack(fill='x', padx=0, pady=0)
        header_frame.pack_propagate(False)
        
        # Título principal
        titulo = tk.Label(header_frame,
                         text="SISTEMA DE DEBIDA DILIGENCIA",
                         font=('Segoe UI', 18, 'bold'),
                         bg='#2C3E50',
                         fg='white')
        titulo.pack(side='left', padx=30, pady=20)
        
        # Subtítulo
        subtitulo = tk.Label(header_frame,
                            text="Gestión Profesional de Cumplimiento",
                            font=('Segoe UI', 10),
                            bg='#2C3E50',
                            fg='#BDC3C7')
        subtitulo.pack(side='left', padx=(0, 30), pady=20)
    
    def cargar_usuarios_guardados(self):
        """Carga los usuarios guardados desde un archivo JSON"""
        try:
            if os.path.exists('usuarios_guardados.json'):
                with open('usuarios_guardados.json', 'r', encoding='utf-8') as f:
                    return json.load(f)
        except:
            pass
        return []
    
    def guardar_usuario(self, usuario):
        """Guarda un usuario en la lista de usuarios guardados"""
        if usuario and usuario not in self.usuarios_guardados:
            self.usuarios_guardados.append(usuario)
            try:
                with open('usuarios_guardados.json', 'w', encoding='utf-8') as f:
                    json.dump(self.usuarios_guardados, f, ensure_ascii=False, indent=2)
            except:
                pass
    
    def cargar_serapios_guardados(self):
        """Carga los serapios guardados desde un archivo JSON"""
        try:
            if os.path.exists('serapios_guardados.json'):
                with open('serapios_guardados.json', 'r', encoding='utf-8') as f:
                    return json.load(f)
        except:
            pass
        return []
    
    def guardar_serapio(self, serapio):
        """Guarda un serapio en la lista de serapios guardados"""
        if serapio and serapio not in self.serapios_guardados:
            self.serapios_guardados.append(serapio)
            try:
                with open('serapios_guardados.json', 'w', encoding='utf-8') as f:
                    json.dump(self.serapios_guardados, f, ensure_ascii=False, indent=2)
            except:
                pass
    
    def cargar_sujetos_guardados(self):
        """Carga los sujetos guardados desde un archivo JSON"""
        try:
            if os.path.exists('sujetos_guardados.json'):
                with open('sujetos_guardados.json', 'r', encoding='utf-8') as f:
                    self.sujetos = json.load(f)
        except:
            pass
    
    def guardar_sujetos_archivo(self):
        """Guarda todos los sujetos en un archivo JSON"""
        try:
            with open('sujetos_guardados.json', 'w', encoding='utf-8') as f:
                json.dump(self.sujetos, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error al guardar sujetos: {e}")
        
    def cargar_configuraciones_default(self):
        """Carga configuraciones desde archivo o crea las predeterminadas"""
        try:
            if os.path.exists('configuraciones.json'):
                with open('configuraciones.json', 'r', encoding='utf-8') as f:
                    return json.load(f)
        except:
            pass
        
        return {
            'tipos_solicitud': ['proveedor', 'cliente', 'socio', 'empleado', 'otro'],
            'fuentes_investigacion': [
                'Ministerio Público',
                'Diarios Hondureños',
                'Listas de restricción',
                'OFAC',
                'Infornet'
            ],
            'nombres_pestanas': {
                'fecha_solicitud': 'Fecha de Solicitud',
                'usuario_requirente': 'Usuario Requirente',
                'tipo_solicitud': 'Tipo de Solicitud',
                'descripcion': 'Descripción',
                'serapio': 'Serapio',
                'fuente_info': 'Fuente de Información'
            },
            'titulo_documento': 'INFORME DE DEBIDA DILIGENCIA'
        }
    
    def guardar_configuraciones_archivo(self):
        """Guarda las configuraciones en un archivo JSON"""
        try:
            with open('configuraciones.json', 'w', encoding='utf-8') as f:
                json.dump(self.configuraciones, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error al guardar configuraciones: {e}")
        
    def cargar_plantilla_default(self):
        return {
            'objetivo': 'Cumplir con lo establecido en la normativa de prevención de lavado de activos',
            'conclusion_template': 'Tras el análisis de las fuentes consultadas y la información disponible públicamente, no se han identificado indicios que sugieran la existencia de riesgos reputacionales o legales relevantes asociados a {nombres}. Por lo tanto, se considera que el nivel de riesgo de esta operación, en relación con el lavado de activos, es **{nivel_riesgo}**.',
            'compromiso': 'El área de Cumplimiento reitera su compromiso en cumplir con todos los requerimientos establecidos en las leyes nacionales y estándares internacionales.',
            'lugar_emision': 'Tegucigalpa, Honduras',
            'jefe_cumplimiento': '',
            'analista': '',
            'incisos_adicionales': []  # Nueva lista para incisos personalizados
        }
    
    def crear_tab_datos_generales(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="📋 Datos Generales")
        
        # Frame principal con fondo blanco
        main_container = tk.Frame(tab, bg='#ECF0F1')
        main_container.pack(fill='both', expand=True)
        
        # Frame principal
        frame = ttk.LabelFrame(main_container, text="Información General del Caso", padding=25)
        frame.pack(fill='both', expand=True, padx=20, pady=20)
        
        # ID de Gestión (nuevo campo al inicio)
        ttk.Label(frame, text="ID de Gestión:", font=('Segoe UI', 10, 'bold')).grid(row=0, column=0, sticky='w', pady=8)
        gestion_frame = tk.Frame(frame, bg='white')
        gestion_frame.grid(row=0, column=1, pady=8, padx=15, sticky='ew')
        self.id_gestion = tk.Entry(gestion_frame, width=30, font=('Segoe UI', 10), relief='solid', borderwidth=1)
        self.id_gestion.pack(side='left', padx=(0, 5))
        ttk.Button(gestion_frame, text="Generar ID", command=self.generar_id_gestion, width=12).pack(side='left')
        
        # Fecha de solicitud
        ttk.Label(frame, text=self.configuraciones['nombres_pestanas']['fecha_solicitud'] + ":", 
                 font=('Segoe UI', 10, 'bold')).grid(row=1, column=0, sticky='w', pady=8)
        self.fecha_solicitud = tk.Entry(frame, width=45, font=('Segoe UI', 10), relief='solid', borderwidth=1)
        self.fecha_solicitud.insert(0, datetime.now().strftime("%d de %B del %Y"))
        self.fecha_solicitud.grid(row=1, column=1, pady=8, padx=15, sticky='ew')
        
        # Usuario requirente con autocompletado
        ttk.Label(frame, text=self.configuraciones['nombres_pestanas']['usuario_requirente'] + ":", 
                 font=('Segoe UI', 10, 'bold')).grid(row=2, column=0, sticky='w', pady=8)
        usuario_frame = tk.Frame(frame, bg='white')
        usuario_frame.grid(row=2, column=1, pady=8, padx=15, sticky='ew')
        self.usuario_requirente = ttk.Combobox(usuario_frame, width=35, values=self.usuarios_guardados, font=('Segoe UI', 10))
        self.usuario_requirente.pack(side='left', fill='x', expand=True, padx=(0, 5))
        ttk.Button(usuario_frame, text="💾", command=self.guardar_usuario_actual, width=3).pack(side='left')
        
        # Tipo de solicitud (solo lectura, se modifica desde configuraciones)
        ttk.Label(frame, text=self.configuraciones['nombres_pestanas']['tipo_solicitud'] + ":", 
                 font=('Segoe UI', 10, 'bold')).grid(row=3, column=0, sticky='w', pady=8)
        self.tipo_solicitud = ttk.Combobox(frame, width=42, values=self.configuraciones['tipos_solicitud'], 
                                          font=('Segoe UI', 10), state='readonly')
        self.tipo_solicitud.set('proveedor')
        self.tipo_solicitud.grid(row=3, column=1, pady=8, padx=15, sticky='ew')
        
        # Descripción
        ttk.Label(frame, text=self.configuraciones['nombres_pestanas']['descripcion'] + ":", 
                 font=('Segoe UI', 10, 'bold')).grid(row=4, column=0, sticky='w', pady=8)
        self.descripcion_solicitud = tk.Entry(frame, width=45, font=('Segoe UI', 10), relief='solid', borderwidth=1)
        self.descripcion_solicitud.insert(0, 'Solicitud de permisos')
        self.descripcion_solicitud.grid(row=4, column=1, pady=8, padx=15, sticky='ew')
        
        # Serapio (NUEVO CAMPO)
        ttk.Label(frame, text=self.configuraciones['nombres_pestanas']['serapio'] + ":", 
                 font=('Segoe UI', 10, 'bold')).grid(row=5, column=0, sticky='w', pady=8)
        serapio_frame = tk.Frame(frame, bg='white')
        serapio_frame.grid(row=5, column=1, pady=8, padx=15, sticky='ew')
        self.serapio = ttk.Combobox(serapio_frame, width=35, values=self.serapios_guardados, font=('Segoe UI', 10))
        self.serapio.pack(side='left', fill='x', expand=True, padx=(0, 5))
        ttk.Button(serapio_frame, text="💾", command=self.guardar_serapio_actual, width=3).pack(side='left')
        
        # Fuente de información
        ttk.Label(frame, text=self.configuraciones['nombres_pestanas']['fuente_info'] + ":", 
                 font=('Segoe UI', 10, 'bold')).grid(row=6, column=0, sticky='nw', pady=8)
        self.fuente_info = scrolledtext.ScrolledText(frame, width=45, height=4, font=('Segoe UI', 9), 
                                                     relief='solid', borderwidth=1, wrap='word')
        self.fuente_info.insert(1.0, 'Búsqueda en medios de comunicación hondureños y bases de datos públicas y privadas.')
        self.fuente_info.grid(row=6, column=1, pady=8, padx=15, sticky='ew')
        
        # Configurar expansión de columnas
        frame.columnconfigure(1, weight=1)
    
    def generar_id_gestion(self):
        """Genera un ID único para la gestión"""
        fecha = datetime.now().strftime("%Y%m%d")
        hora = datetime.now().strftime("%H%M%S")
        id_gestion = f"DD-{fecha}-{hora}"
        self.id_gestion.delete(0, tk.END)
        self.id_gestion.insert(0, id_gestion)
        self.gestion_actual = id_gestion
        messagebox.showinfo("ID Generado", f"ID de Gestión: {id_gestion}\n\nTodos los sujetos agregados a partir de ahora pertenecerán a esta gestión.")
    
    def guardar_usuario_actual(self):
        """Guarda el usuario actual en la lista"""
        usuario = self.usuario_requirente.get().strip()
        if usuario:
            self.guardar_usuario(usuario)
            self.usuario_requirente['values'] = self.usuarios_guardados
            messagebox.showinfo("Guardado", f"Usuario '{usuario}' guardado correctamente")
    
    def guardar_serapio_actual(self):
        """Guarda el serapio actual en la lista"""
        serapio = self.serapio.get().strip()
        if serapio:
            self.guardar_serapio(serapio)
            self.serapio['values'] = self.serapios_guardados
            messagebox.showinfo("Guardado", f"Serapio '{serapio}' guardado correctamente")
    
    def crear_tab_sujetos(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="👤 Sujetos de Investigación")
        
        # Configurar fondo
        main_container = tk.Frame(tab, bg='#ECF0F1')
        main_container.pack(fill='both', expand=True)
        
        # Frame de búsqueda (NUEVO)
        frame_busqueda = ttk.LabelFrame(main_container, text="🔍 Buscar Sujetos", padding=15)
        frame_busqueda.pack(fill='x', padx=20, pady=(20, 10))
        
        busqueda_frame = tk.Frame(frame_busqueda, bg='white')
        busqueda_frame.pack(fill='x')
        
        ttk.Label(busqueda_frame, text="Buscar por:", font=('Segoe UI', 10, 'bold')).pack(side='left', padx=5)
        
        self.tipo_busqueda = ttk.Combobox(busqueda_frame, width=15, values=['Nombre', 'Identificación', 'ID Gestión'], 
                                         font=('Segoe UI', 10), state='readonly')
        self.tipo_busqueda.set('Nombre')
        self.tipo_busqueda.pack(side='left', padx=5)
        
        self.campo_busqueda = tk.Entry(busqueda_frame, width=30, font=('Segoe UI', 10), relief='solid', borderwidth=1)
        self.campo_busqueda.pack(side='left', padx=5)
        self.campo_busqueda.bind('<KeyRelease>', self.buscar_sujeto_en_tiempo_real)
        
        ttk.Button(busqueda_frame, text="🔍 Buscar", command=self.buscar_sujeto).pack(side='left', padx=5)
        ttk.Button(busqueda_frame, text="🔄 Mostrar Todos", command=self.mostrar_todos_sujetos).pack(side='left', padx=5)
        
        # Frame de entrada
        frame_entrada = ttk.LabelFrame(main_container, text="Agregar Nuevo Sujeto", padding=25)
        frame_entrada.pack(fill='x', padx=20, pady=10)
        
        # Campos en una sola fila para nombre e identificación (más angostos)
        campos_frame = tk.Frame(frame_entrada, bg='white')
        campos_frame.grid(row=0, column=0, columnspan=2, sticky='ew', pady=8)
        
        ttk.Label(campos_frame, text="Nombre:", font=('Segoe UI', 10, 'bold')).pack(side='left', padx=(0, 5))
        self.nombre_sujeto = tk.Entry(campos_frame, width=25, font=('Segoe UI', 10), relief='solid', borderwidth=1)
        self.nombre_sujeto.pack(side='left', padx=5)
        # Validación: solo letras y espacios
        vcmd_nombre = (self.root.register(self.validar_nombre), '%P')
        self.nombre_sujeto.config(validate='key', validatecommand=vcmd_nombre)
        
        ttk.Label(campos_frame, text="Identificación:", font=('Segoe UI', 10, 'bold')).pack(side='left', padx=(15, 5))
        self.identificacion_sujeto = tk.Entry(campos_frame, width=20, font=('Segoe UI', 10), relief='solid', borderwidth=1)
        self.identificacion_sujeto.pack(side='left', padx=5)
        # Validación: solo números y guiones
        vcmd_id = (self.root.register(self.validar_identificacion), '%P')
        self.identificacion_sujeto.config(validate='key', validatecommand=vcmd_id)
        
        # Descripción en segunda fila (ancho completo)
        ttk.Label(frame_entrada, text="Descripción:", font=('Segoe UI', 10, 'bold')).grid(row=1, column=0, sticky='w', pady=8)
        self.descripcion_sujeto = tk.Entry(frame_entrada, width=60, font=('Segoe UI', 10), relief='solid', borderwidth=1)
        self.descripcion_sujeto.grid(row=1, column=1, pady=8, padx=15, sticky='ew')
        
        # Botón de agregar con estilo
        btn_frame = tk.Frame(frame_entrada, bg='white')
        btn_frame.grid(row=2, column=0, columnspan=2, pady=15)
        ttk.Button(btn_frame, text="➕ Agregar Sujeto", command=self.agregar_sujeto).pack()
        
        frame_entrada.columnconfigure(1, weight=1)
        
        # Frame de lista
        frame_lista = ttk.LabelFrame(main_container, text="Sujetos Registrados", padding=25)
        frame_lista.pack(fill='both', expand=True, padx=20, pady=(10, 20))
        
        # Treeview con estilo mejorado (incluye columna de ID de Gestión y columna de Acciones)
        tree_frame = tk.Frame(frame_lista, bg='white')
        tree_frame.pack(fill='both', expand=True)
        
        self.tree_sujetos = ttk.Treeview(tree_frame, columns=('Nombre', 'Identificación', 'Descripción', 'Gestión', 'Acciones'), 
                                         show='headings', height=12)
        self.tree_sujetos.heading('Nombre', text='Nombre Completo')
        self.tree_sujetos.heading('Identificación', text='Identificación')
        self.tree_sujetos.heading('Descripción', text='Descripción')
        self.tree_sujetos.heading('Gestión', text='ID de Gestión')
        self.tree_sujetos.heading('Acciones', text='Eliminar')
        
        # Ajustar anchos
        self.tree_sujetos.column('Nombre', width=180)
        self.tree_sujetos.column('Identificación', width=110)
        self.tree_sujetos.column('Descripción', width=180)
        self.tree_sujetos.column('Gestión', width=130)
        self.tree_sujetos.column('Acciones', width=70)
        
        # Bind para detectar clicks en la columna de acciones
        self.tree_sujetos.bind('<Button-1>', self.on_tree_click)
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(tree_frame, orient='vertical', command=self.tree_sujetos.yview)
        self.tree_sujetos.configure(yscrollcommand=scrollbar.set)
        
        self.tree_sujetos.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        # Botones (sin botón eliminar individual ya que está en cada fila)
        btn_frame2 = tk.Frame(frame_lista, bg='white')
        btn_frame2.pack(pady=10)
        ttk.Button(btn_frame2, text="📋 Cargar Seleccionado", command=self.cargar_sujeto_seleccionado).pack(side='left', padx=5)
        ttk.Button(btn_frame2, text="🔍 Filtrar por Gestión", command=self.filtrar_por_gestion).pack(side='left', padx=5)
    
    def validar_nombre(self, texto):
        """Valida que solo se ingresen letras, espacios y algunos caracteres especiales en nombres"""
        if texto == "":
            return True
        # Permitir letras (incluyendo acentos y ñ), espacios, puntos, comas y apóstrofes
        import re
        return bool(re.match(r"^[a-zA-ZáéíóúÁÉÍÓÚñÑüÜ\s.,'-]+$", texto))
    
    def validar_identificacion(self, texto):
        """Valida que solo se ingresen números y guiones en identificación"""
        if texto == "":
            return True
        # Solo números y guiones
        import re
        return bool(re.match(r"^[0-9-]+$", texto))
    
    def filtrar_por_gestion(self):
        """Muestra solo los sujetos de una gestión específica"""
        gestion_id = self.id_gestion.get().strip()
        if not gestion_id:
            messagebox.showwarning("Advertencia", "Primero debe ingresar o generar un ID de Gestión")
            return
        
        sujetos_filtrados = [s for s in self.sujetos if s.get('id_gestion') == gestion_id]
        
        if not sujetos_filtrados:
            messagebox.showinfo("Sin resultados", f"No hay sujetos asociados a la gestión: {gestion_id}")
            return
        
        mensaje = f"Sujetos de la gestión {gestion_id}:\n\n"
        for i, sujeto in enumerate(sujetos_filtrados, 1):
            mensaje += f"{i}. {sujeto['nombre']} ({sujeto['identificacion']})\n"
        
        messagebox.showinfo("Sujetos Filtrados", mensaje)
    
    def buscar_sujeto(self):
        """Busca sujetos según el criterio seleccionado"""
        tipo = self.tipo_busqueda.get()
        termino = self.campo_busqueda.get().strip().lower()
        
        if not termino:
            messagebox.showwarning("Advertencia", "Ingrese un término de búsqueda")
            return
        
        # Limpiar el árbol
        for item in self.tree_sujetos.get_children():
            self.tree_sujetos.delete(item)
        
        # Buscar y mostrar resultados
        encontrados = 0
        for sujeto in self.sujetos:
            mostrar = False
            
            if tipo == 'Nombre' and termino in sujeto['nombre'].lower():
                mostrar = True
            elif tipo == 'Identificación' and termino in sujeto['identificacion'].lower():
                mostrar = True
            elif tipo == 'ID Gestión' and termino in sujeto.get('id_gestion', '').lower():
                mostrar = True
            
            if mostrar:
                self.tree_sujetos.insert('', 'end', values=(
                    sujeto['nombre'],
                    sujeto['identificacion'],
                    sujeto['descripcion'],
                    sujeto.get('id_gestion', 'N/A'),
                    '🗑️'
                ))
                encontrados += 1
        
        if encontrados == 0:
            messagebox.showinfo("Sin resultados", f"No se encontraron sujetos con '{termino}' en {tipo}")
    
    def buscar_sujeto_en_tiempo_real(self, event=None):
        """Busca sujetos mientras el usuario escribe"""
        tipo = self.tipo_busqueda.get()
        termino = self.campo_busqueda.get().strip().lower()
        
        # Limpiar el árbol
        for item in self.tree_sujetos.get_children():
            self.tree_sujetos.delete(item)
        
        # Si no hay término, mostrar todos
        if not termino:
            self.mostrar_todos_sujetos()
            return
        
        # Buscar y mostrar resultados
        for sujeto in self.sujetos:
            mostrar = False
            
            if tipo == 'Nombre' and termino in sujeto['nombre'].lower():
                mostrar = True
            elif tipo == 'Identificación' and termino in sujeto['identificacion'].lower():
                mostrar = True
            elif tipo == 'ID Gestión' and termino in sujeto.get('id_gestion', '').lower():
                mostrar = True
            
            if mostrar:
                self.tree_sujetos.insert('', 'end', values=(
                    sujeto['nombre'],
                    sujeto['identificacion'],
                    sujeto['descripcion'],
                    sujeto.get('id_gestion', 'N/A'),
                    '🗑️'
                ))
    
    def mostrar_todos_sujetos(self):
        """Muestra todos los sujetos en el árbol"""
        # Limpiar el árbol
        for item in self.tree_sujetos.get_children():
            self.tree_sujetos.delete(item)
        
        # Agregar todos los sujetos
        for sujeto in self.sujetos:
            self.tree_sujetos.insert('', 'end', values=(
                sujeto['nombre'],
                sujeto['identificacion'],
                sujeto['descripcion'],
                sujeto.get('id_gestion', 'N/A'),
                '🗑️'
            ))
        
        # Limpiar campo de búsqueda
        self.campo_busqueda.delete(0, tk.END)
    
    def cargar_sujeto_seleccionado(self):
        """Carga los datos del sujeto seleccionado en los campos de entrada"""
        selected = self.tree_sujetos.selection()
        if not selected:
            messagebox.showwarning("Advertencia", "Seleccione un sujeto para cargar")
            return
        
        # Obtener índice del sujeto seleccionado
        item = self.tree_sujetos.item(selected[0])
        valores = item['values']
        
        # Buscar el sujeto completo en la lista
        for sujeto in self.sujetos:
            if (sujeto['nombre'] == valores[0] and 
                sujeto['identificacion'] == valores[1] and
                sujeto.get('id_gestion', 'N/A') == valores[3]):
                
                # Cargar datos en los campos
                self.nombre_sujeto.delete(0, tk.END)
                self.nombre_sujeto.insert(0, sujeto['nombre'])
                
                self.identificacion_sujeto.delete(0, tk.END)
                self.identificacion_sujeto.insert(0, sujeto['identificacion'])
                
                self.descripcion_sujeto.delete(0, tk.END)
                self.descripcion_sujeto.insert(0, sujeto['descripcion'])
                
                messagebox.showinfo("Cargado", f"Datos de '{sujeto['nombre']}' cargados en los campos")
                break
    
    def eliminar_todos_sujetos(self):
        """Elimina todos los sujetos guardados sin confirmación"""
        if not self.sujetos:
            messagebox.showinfo("Información", "No hay sujetos para eliminar")
            return
        
        # Limpiar lista de sujetos
        self.sujetos.clear()
        
        # Limpiar árbol
        for item in self.tree_sujetos.get_children():
            self.tree_sujetos.delete(item)
        
        # Guardar cambios (archivo vacío)
        self.guardar_sujetos_archivo()
        
        messagebox.showinfo("Éxito", "Todos los sujetos han sido eliminados")
    
    def crear_tab_resultados(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="🔍 Resultados de Investigación")
        
        main_container = tk.Frame(tab, bg='#ECF0F1')
        main_container.pack(fill='both', expand=True)
        
        frame = ttk.LabelFrame(main_container, text="Fuentes Consultadas", padding=20)
        frame.pack(fill='both', expand=True, padx=20, pady=20)
        
        # Crear campos para cada fuente
        self.resultados = {}
        
        for i, fuente in enumerate(self.configuraciones['fuentes_investigacion']):
            ttk.Label(frame, text=f"{fuente}:", font=('Segoe UI', 10, 'bold')).grid(row=i*2, column=0, sticky='w', pady=5)
            
            resultado = tk.Entry(frame, width=60, font=('Segoe UI', 10))
            resultado.insert(0, 'No se encontraron referencias.')
            resultado.grid(row=i*2+1, column=0, pady=2, padx=20, sticky='ew')
            
            self.resultados[fuente] = resultado
        
        # Nivel de riesgo
        ttk.Label(frame, text="Nivel de Riesgo:", font=('Segoe UI', 10, 'bold')).grid(
            row=len(self.configuraciones['fuentes_investigacion'])*2, column=0, sticky='w', pady=10)
        self.nivel_riesgo = ttk.Combobox(frame, width=20, values=['bajo', 'medio', 'alto'], font=('Segoe UI', 10))
        self.nivel_riesgo.set('bajo')
        self.nivel_riesgo.grid(row=len(self.configuraciones['fuentes_investigacion'])*2+1, column=0, 
                              pady=2, padx=20, sticky='w')
        
        frame.columnconfigure(0, weight=1)
    
    def crear_tab_plantilla(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="📝 Modificar Plantilla")
        
        main_container = tk.Frame(tab, bg='#ECF0F1')
        main_container.pack(fill='both', expand=True)
        
        # Frame principal con scroll
        canvas = tk.Canvas(main_container, bg='#ECF0F1')
        scrollbar = ttk.Scrollbar(main_container, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='#ECF0F1')
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        frame = ttk.LabelFrame(scrollable_frame, text="Personalizar Plantilla", padding=20)
        frame.pack(fill='both', expand=True, padx=20, pady=20)
        
        # Objetivo
        ttk.Label(frame, text="Objetivo del Informe:", font=('Segoe UI', 10, 'bold')).pack(anchor='w', pady=5)
        self.objetivo_text = scrolledtext.ScrolledText(frame, width=80, height=3, font=('Segoe UI', 10))
        self.objetivo_text.insert(1.0, self.plantilla_modificable['objetivo'])
        self.objetivo_text.pack(pady=5, fill='x')
        
        # Conclusión template
        ttk.Label(frame, text="Plantilla de Conclusión (use {nombres} y {nivel_riesgo}):", 
                 font=('Segoe UI', 10, 'bold')).pack(anchor='w', pady=5)
        self.conclusion_text = scrolledtext.ScrolledText(frame, width=80, height=4, font=('Segoe UI', 10))
        self.conclusion_text.insert(1.0, self.plantilla_modificable['conclusion_template'])
        self.conclusion_text.pack(pady=5, fill='x')
        
        # Compromiso
        ttk.Label(frame, text="Texto de Compromiso:", font=('Segoe UI', 10, 'bold')).pack(anchor='w', pady=5)
        self.compromiso_text = scrolledtext.ScrolledText(frame, width=80, height=3, font=('Segoe UI', 10))
        self.compromiso_text.insert(1.0, self.plantilla_modificable['compromiso'])
        self.compromiso_text.pack(pady=5, fill='x')
        
        # NUEVA SECCIÓN: Incisos Adicionales
        ttk.Separator(frame, orient='horizontal').pack(fill='x', pady=20)
        
        inciso_header = tk.Frame(frame, bg='white')
        inciso_header.pack(fill='x', pady=10)
        ttk.Label(inciso_header, text="Incisos Adicionales para la Conclusión:", 
                 font=('Segoe UI', 11, 'bold')).pack(side='left')
        ttk.Button(inciso_header, text="➕ Agregar Inciso", command=self.agregar_inciso).pack(side='right')
        
        # Lista de incisos
        self.frame_incisos = tk.Frame(frame, bg='white')
        self.frame_incisos.pack(fill='both', expand=True, pady=10)
        
        self.incisos_widgets = []
        self.cargar_incisos_guardados()
        
        ttk.Separator(frame, orient='horizontal').pack(fill='x', pady=20)
        
        # Datos de firma
        frame_firma = ttk.Frame(frame)
        frame_firma.pack(fill='x', pady=10)
        
        ttk.Label(frame_firma, text="Lugar de Emisión:", font=('Segoe UI', 10)).grid(row=0, column=0, sticky='w', pady=5)
        self.lugar_emision = tk.Entry(frame_firma, width=40, font=('Segoe UI', 10))
        self.lugar_emision.insert(0, self.plantilla_modificable['lugar_emision'])
        self.lugar_emision.grid(row=0, column=1, pady=5, padx=10, sticky='ew')
        
        ttk.Label(frame_firma, text="Jefe de Cumplimiento:", font=('Segoe UI', 10)).grid(row=1, column=0, sticky='w', pady=5)
        self.jefe_cumplimiento = tk.Entry(frame_firma, width=40, font=('Segoe UI', 10))
        self.jefe_cumplimiento.grid(row=1, column=1, pady=5, padx=10, sticky='ew')
        
        ttk.Label(frame_firma, text="Analista:", font=('Segoe UI', 10)).grid(row=2, column=0, sticky='w', pady=5)
        self.analista = tk.Entry(frame_firma, width=40, font=('Segoe UI', 10))
        self.analista.grid(row=2, column=1, pady=5, padx=10, sticky='ew')
        
        frame_firma.columnconfigure(1, weight=1)
        
        ttk.Button(frame, text="💾 Guardar Cambios de Plantilla", command=self.guardar_plantilla).pack(pady=20)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
    
    def agregar_inciso(self):
        """Agrega un nuevo campo para inciso personalizado"""
        inciso_frame = tk.Frame(self.frame_incisos, bg='white', relief='ridge', borderwidth=1)
        inciso_frame.pack(fill='x', pady=5, padx=10)
        
        letra_frame = tk.Frame(inciso_frame, bg='white')
        letra_frame.pack(side='left', padx=5, pady=5)
        
        numero_inciso = len(self.incisos_widgets) + 3  # Empieza en 'c' (después de a y b)
        letra = chr(96 + numero_inciso)  # 'c', 'd', 'e', etc.
        
        ttk.Label(letra_frame, text=f"{letra}.", font=('Segoe UI', 10, 'bold')).pack()
        
        text_widget = scrolledtext.ScrolledText(inciso_frame, width=70, height=3, font=('Segoe UI', 9), wrap='word')
        text_widget.pack(side='left', fill='both', expand=True, padx=5, pady=5)
        
        btn_eliminar = ttk.Button(inciso_frame, text="🗑️", width=3, 
                                  command=lambda: self.eliminar_inciso(inciso_frame, text_widget))
        btn_eliminar.pack(side='right', padx=5, pady=5)
        
        self.incisos_widgets.append(text_widget)
    
    def eliminar_inciso(self, frame, widget):
        """Elimina un inciso personalizado"""
        if messagebox.askyesno("Confirmar", "¿Desea eliminar este inciso?"):
            self.incisos_widgets.remove(widget)
            frame.destroy()
            self.renumerar_incisos()
    
    def renumerar_incisos(self):
        """Renumera los incisos después de eliminar uno"""
        for i, widget_frame in enumerate(self.frame_incisos.winfo_children()):
            numero_inciso = i + 3
            letra = chr(96 + numero_inciso)
            # Actualizar la letra del label
            for child in widget_frame.winfo_children():
                if isinstance(child, tk.Frame):
                    for label in child.winfo_children():
                        if isinstance(label, ttk.Label):
                            label.config(text=f"{letra}.")
    
    def cargar_incisos_guardados(self):
        """Carga los incisos guardados en la plantilla"""
        if 'incisos_adicionales' in self.plantilla_modificable:
            for texto_inciso in self.plantilla_modificable['incisos_adicionales']:
                self.agregar_inciso()
                self.incisos_widgets[-1].insert(1.0, texto_inciso)
    
    def crear_tab_configuraciones(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="⚙️ Configuraciones")
        
        main_container = tk.Frame(tab, bg='#ECF0F1')
        main_container.pack(fill='both', expand=True)
        
        # Frame con scroll
        canvas = tk.Canvas(main_container, bg='#ECF0F1')
        scrollbar = ttk.Scrollbar(main_container, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='#ECF0F1')
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        frame = ttk.LabelFrame(scrollable_frame, text="Personalizar Opciones del Sistema", padding=20)
        frame.pack(fill='both', expand=True, padx=20, pady=20)
        
        # Tipos de Solicitud
        ttk.Label(frame, text="Tipos de Solicitud:", font=('Segoe UI', 11, 'bold')).pack(anchor='w', pady=10)
        ttk.Label(frame, text="Ingrese los tipos de solicitud separados por comas:", 
                 font=('Segoe UI', 9)).pack(anchor='w', pady=2)
        
        self.tipos_solicitud_text = tk.Entry(frame, width=80, font=('Segoe UI', 10))
        self.tipos_solicitud_text.insert(0, ', '.join(self.configuraciones['tipos_solicitud']))
        self.tipos_solicitud_text.pack(pady=5, fill='x')
        
        ttk.Separator(frame, orient='horizontal').pack(fill='x', pady=20)
        
        # Fuentes de Investigación
        ttk.Label(frame, text="Fuentes de Investigación:", font=('Segoe UI', 11, 'bold')).pack(anchor='w', pady=10)
        ttk.Label(frame, text="Ingrese las fuentes de investigación (una por línea):", 
                 font=('Segoe UI', 9)).pack(anchor='w', pady=2)
        
        self.fuentes_text = scrolledtext.ScrolledText(frame, width=80, height=8, font=('Segoe UI', 10))
        self.fuentes_text.insert(1.0, '\n'.join(self.configuraciones['fuentes_investigacion']))
        self.fuentes_text.pack(pady=5, fill='x')
        
        ttk.Separator(frame, orient='horizontal').pack(fill='x', pady=20)
        
        # NUEVA SECCIÓN: Nombres de Pestañas
        ttk.Label(frame, text="Nombres de Campos en Información General:", 
                 font=('Segoe UI', 11, 'bold')).pack(anchor='w', pady=10)
        ttk.Label(frame, text="Personalice los nombres de los campos que aparecen en Datos Generales:", 
                 font=('Segoe UI', 9)).pack(anchor='w', pady=2)
        
        nombres_frame = tk.Frame(frame, bg='white', relief='groove', borderwidth=2)
        nombres_frame.pack(fill='x', pady=10, padx=5)
        
        self.campos_nombres = {}
        campos_config = [
            ('fecha_solicitud', 'Fecha de Solicitud'),
            ('usuario_requirente', 'Usuario Requirente'),
            ('tipo_solicitud', 'Tipo de Solicitud'),
            ('descripcion', 'Descripción'),
            ('serapio', 'Serapio'),
            ('fuente_info', 'Fuente de Información')
        ]
        
        for i, (clave, nombre_default) in enumerate(campos_config):
            campo_frame = tk.Frame(nombres_frame, bg='white')
            campo_frame.pack(fill='x', padx=10, pady=5)
            
            ttk.Label(campo_frame, text=f"{nombre_default}:", 
                     font=('Segoe UI', 9, 'bold'), width=25).pack(side='left')
            
            entry = tk.Entry(campo_frame, width=40, font=('Segoe UI', 9))
            valor_actual = self.configuraciones['nombres_pestanas'].get(clave, nombre_default)
            entry.insert(0, valor_actual)
            entry.pack(side='left', padx=5)
            
            self.campos_nombres[clave] = entry
        
        ttk.Separator(frame, orient='horizontal').pack(fill='x', pady=20)
        
        # NUEVA SECCIÓN: Título del Documento
        ttk.Label(frame, text="Título del Documento Word:", 
                 font=('Segoe UI', 11, 'bold')).pack(anchor='w', pady=10)
        ttk.Label(frame, text="Personalice el título principal que aparece en el documento generado:", 
                 font=('Segoe UI', 9)).pack(anchor='w', pady=2)
        
        self.titulo_documento = tk.Entry(frame, width=80, font=('Segoe UI', 10))
        titulo_actual = self.configuraciones.get('titulo_documento', 'INFORME DE DEBIDA DILIGENCIA')
        self.titulo_documento.insert(0, titulo_actual)
        self.titulo_documento.pack(pady=5, fill='x')
        
        ttk.Label(frame, text="Nota: Después de guardar, los cambios se aplicarán al reiniciar la pestaña", 
                 font=('Segoe UI', 9, 'italic'), foreground='#E74C3C').pack(pady=10)
        
        frame_botones = tk.Frame(frame, bg='white')
        frame_botones.pack(pady=20)
        
        ttk.Button(frame_botones, text="💾 Guardar Configuraciones", 
                  command=self.guardar_configuraciones).pack(side='left', padx=5)
        ttk.Button(frame_botones, text="🔄 Restaurar Predeterminados", 
                  command=self.restaurar_configuraciones).pack(side='left', padx=5)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
    
    def guardar_configuraciones(self):
        try:
            # Actualizar tipos de solicitud
            tipos_texto = self.tipos_solicitud_text.get().strip()
            self.configuraciones['tipos_solicitud'] = [t.strip() for t in tipos_texto.split(',') if t.strip()]
            
            # Actualizar fuentes
            fuentes_texto = self.fuentes_text.get(1.0, tk.END).strip()
            self.configuraciones['fuentes_investigacion'] = [f.strip() for f in fuentes_texto.split('\n') if f.strip()]
            
            # Actualizar nombres de pestañas
            for clave, entry in self.campos_nombres.items():
                self.configuraciones['nombres_pestanas'][clave] = entry.get().strip()
            
            # Actualizar título del documento
            self.configuraciones['titulo_documento'] = self.titulo_documento.get().strip()
            
            # Guardar en archivo
            self.guardar_configuraciones_archivo()
            
            # Actualizar el combobox de tipo de solicitud
            self.tipo_solicitud['values'] = self.configuraciones['tipos_solicitud']
            
            # Actualizar la pestaña de resultados
            self.actualizar_tab_resultados()
            
            messagebox.showinfo("Éxito", "Configuraciones guardadas correctamente.\n\n" +
                              "La pestaña de Resultados ha sido actualizada.\n" +
                              "Los nombres de campos se actualizarán al reiniciar la aplicación.")
        except Exception as e:
            messagebox.showerror("Error", f"Error al guardar configuraciones: {str(e)}")
    
    def restaurar_configuraciones(self):
        if messagebox.askyesno("Confirmar", "¿Desea restaurar las configuraciones predeterminadas?"):
            # Restaurar sin el archivo
            config_temp = self.cargar_configuraciones_default()
            # Mantener nombres_pestanas si no existen en default
            if 'nombres_pestanas' not in config_temp:
                config_temp['nombres_pestanas'] = {
                    'fecha_solicitud': 'Fecha de Solicitud',
                    'usuario_requirente': 'Usuario Requirente',
                    'tipo_solicitud': 'Tipo de Solicitud',
                    'descripcion': 'Descripción',
                    'serapio': 'Serapio',
                    'fuente_info': 'Fuente de Información'
                }
            
            self.configuraciones = config_temp
            
            # Actualizar campos en la interfaz
            self.tipos_solicitud_text.delete(0, tk.END)
            self.tipos_solicitud_text.insert(0, ', '.join(self.configuraciones['tipos_solicitud']))
            
            self.fuentes_text.delete(1.0, tk.END)
            self.fuentes_text.insert(1.0, '\n'.join(self.configuraciones['fuentes_investigacion']))
            
            for clave, entry in self.campos_nombres.items():
                entry.delete(0, tk.END)
                entry.insert(0, self.configuraciones['nombres_pestanas'][clave])
            
            # Actualizar título del documento
            self.titulo_documento.delete(0, tk.END)
            titulo_restaurado = self.configuraciones.get('titulo_documento', 'INFORME DE DEBIDA DILIGENCIA')
            self.titulo_documento.insert(0, titulo_restaurado)
            
            messagebox.showinfo("Éxito", "Configuraciones restauradas. " +
                              "Presione 'Guardar Configuraciones' para aplicar los cambios.")
    
    def actualizar_tab_resultados(self):
        # Guardar valores actuales
        valores_actuales = {}
        for fuente, widget in self.resultados.items():
            valores_actuales[fuente] = widget.get()
        
        # Limpiar resultados actuales
        self.resultados.clear()
        
        # Recrear la pestaña de resultados
        for i in range(self.notebook.index('end')):
            if '🔍' in self.notebook.tab(i, 'text') or 'Resultados' in self.notebook.tab(i, 'text'):
                tab_index = i
                break
        
        # Eliminar la pestaña antigua
        self.notebook.forget(tab_index)
        
        # Recrear la pestaña
        self.crear_tab_resultados()
        
        # Restaurar valores
        for fuente, widget in self.resultados.items():
            if fuente in valores_actuales:
                widget.delete(0, tk.END)
                widget.insert(0, valores_actuales[fuente])
    
    def crear_tab_generar(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="📄 Generar Documentos")
        
        # Configurar fondo
        main_container = tk.Frame(tab, bg='#ECF0F1')
        main_container.pack(fill='both', expand=True)
        
        frame = ttk.LabelFrame(main_container, text="Opciones de Generación de Documentos", padding=30)
        frame.pack(fill='both', expand=True, padx=20, pady=20)
        
        # Título
        titulo = tk.Label(frame, 
                         text="Seleccione el formato de documento que desea generar:",
                         font=('Segoe UI', 12, 'bold'),
                         bg='white',
                         fg='#2C3E50')
        titulo.pack(pady=(10, 30))
        
        # Botones con descripciones
        btn_container = tk.Frame(frame, bg='white')
        btn_container.pack(pady=10)
        
        # Botón Word
        word_frame = tk.Frame(btn_container, bg='white')
        word_frame.pack(pady=8)
        ttk.Button(word_frame, text="📝 Generar Documento Word (.docx)", 
                  command=self.generar_word, width=50).pack()
        desc_word = tk.Label(word_frame, 
                           text="Genera un informe completo en formato Word con normas APA",
                           font=('Segoe UI', 9, 'italic'),
                           bg='white',
                           fg='#7F8C8D')
        desc_word.pack()
        
        # Botón Excel Completo
        excel_frame = tk.Frame(btn_container, bg='white')
        excel_frame.pack(pady=8)
        ttk.Button(excel_frame, text="📊 Generar Excel Completo", 
                  command=self.generar_excel_completo, width=50).pack()
        desc_excel = tk.Label(excel_frame,
                             text="Base de datos horizontal con todas las debidas diligencias",
                             font=('Segoe UI', 9, 'italic'),
                             bg='white',
                             fg='#7F8C8D')
        desc_excel.pack()
        
        # Botón Excel Pestañas
        pestanas_frame = tk.Frame(btn_container, bg='white')
        pestanas_frame.pack(pady=8)
        ttk.Button(pestanas_frame, text="📑 Generar Excel con Pestañas por Sujeto", 
                  command=self.generar_excel_pestanas, width=50).pack()
        desc_pestanas = tk.Label(pestanas_frame,
                                text="Una pestaña individual para cada sujeto investigado",
                                font=('Segoe UI', 9, 'italic'),
                                bg='white',
                                fg='#7F8C8D')
        desc_pestanas.pack()
        
        # Botón Todo
        todo_frame = tk.Frame(btn_container, bg='white')
        todo_frame.pack(pady=8)
        ttk.Button(todo_frame, text="📦 Generar Todo (Word + Excels)", 
                  command=self.generar_todo, width=50).pack()
        desc_todo = tk.Label(todo_frame,
                           text="Genera todos los formatos disponibles simultáneamente",
                           font=('Segoe UI', 9, 'italic'),
                           bg='white',
                           fg='#7F8C8D')
        desc_todo.pack()
        
        # Separador
        separator = ttk.Separator(frame, orient='horizontal')
        separator.pack(fill='x', pady=20)
        
        # Label de estado con mejor diseño
        status_frame = tk.Frame(frame, bg='#E8F8F5', relief='ridge', borderwidth=2)
        status_frame.pack(fill='x', pady=10)
        
        self.status_label = tk.Label(status_frame, 
                                     text="",
                                     font=('Segoe UI', 10),
                                     bg='#E8F8F5',
                                     fg='#27AE60',
                                     pady=10)
        self.status_label.pack()
    
    def agregar_sujeto(self):
        nombre = self.nombre_sujeto.get().strip()
        identificacion = self.identificacion_sujeto.get().strip()
        descripcion = self.descripcion_sujeto.get().strip()
        id_gestion = self.id_gestion.get().strip()
        
        if not nombre or not identificacion:
            messagebox.showwarning("Advertencia", "Debe ingresar al menos nombre e identificación")
            return
        
        if not id_gestion:
            respuesta = messagebox.askyesno("ID de Gestión", 
                                           "No ha ingresado un ID de Gestión.\n\n" +
                                           "¿Desea generar uno automáticamente?")
            if respuesta:
                self.generar_id_gestion()
                id_gestion = self.id_gestion.get().strip()
            else:
                return
        
        sujeto = {
            'nombre': nombre,
            'identificacion': identificacion,
            'descripcion': descripcion,
            'id_gestion': id_gestion
        }
        
        self.sujetos.append(sujeto)
        self.tree_sujetos.insert('', 'end', values=(nombre, identificacion, descripcion, id_gestion, '🗑️'))
        
        # Guardar en archivo
        self.guardar_sujetos_archivo()
        
        # Limpiar campos
        self.nombre_sujeto.delete(0, tk.END)
        self.identificacion_sujeto.delete(0, tk.END)
        self.descripcion_sujeto.delete(0, tk.END)
        
        messagebox.showinfo("Éxito", f"Sujeto agregado y guardado correctamente\nGestión: {id_gestion}")
    
    def on_tree_click(self, event):
        """Detecta clicks en el árbol y verifica si fue en la columna de eliminar"""
        region = self.tree_sujetos.identify("region", event.x, event.y)
        if region == "cell":
            column = self.tree_sujetos.identify_column(event.x)
            item = self.tree_sujetos.identify_row(event.y)
            
            # Si se hizo click en la columna de "Acciones" (columna #5)
            if column == '#5' and item:
                self.eliminar_sujeto_directo(item)
    
    def eliminar_sujeto_directo(self, item):
        """Elimina un sujeto directamente sin confirmación"""
        # Obtener valores del item
        valores = self.tree_sujetos.item(item)['values']
        
        # Buscar y eliminar de la lista
        for i, sujeto in enumerate(self.sujetos):
            if (sujeto['nombre'] == valores[0] and 
                sujeto['identificacion'] == valores[1] and
                sujeto.get('id_gestion', 'N/A') == valores[3]):
                self.sujetos.pop(i)
                break
        
        # Eliminar del árbol
        self.tree_sujetos.delete(item)
        
        # Guardar cambios en archivo
        self.guardar_sujetos_archivo()
    
    def eliminar_sujeto(self):
        """Elimina el sujeto seleccionado (método alternativo)"""
        selected = self.tree_sujetos.selection()
        if not selected:
            messagebox.showwarning("Advertencia", "Seleccione un sujeto para eliminar")
            return
        
        self.eliminar_sujeto_directo(selected[0])
    
    def guardar_plantilla(self):
        self.plantilla_modificable['objetivo'] = self.objetivo_text.get(1.0, tk.END).strip()
        self.plantilla_modificable['conclusion_template'] = self.conclusion_text.get(1.0, tk.END).strip()
        self.plantilla_modificable['compromiso'] = self.compromiso_text.get(1.0, tk.END).strip()
        self.plantilla_modificable['lugar_emision'] = self.lugar_emision.get().strip()
        self.plantilla_modificable['jefe_cumplimiento'] = self.jefe_cumplimiento.get().strip()
        self.plantilla_modificable['analista'] = self.analista.get().strip()
        
        # Guardar incisos adicionales
        self.plantilla_modificable['incisos_adicionales'] = []
        for widget in self.incisos_widgets:
            texto = widget.get(1.0, tk.END).strip()
            if texto:
                self.plantilla_modificable['incisos_adicionales'].append(texto)
        
        messagebox.showinfo("Éxito", f"Plantilla actualizada correctamente\n\n" +
                          f"Incisos adicionales guardados: {len(self.plantilla_modificable['incisos_adicionales'])}")
    
    def aplicar_formato_apa(self, doc):
        """Aplica formato APA al documento"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        for paragraph in doc.paragraphs:
            # Aplicar fuente a todos los runs
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            
            # Solo aplicar doble espacio si no tiene configuración específica
            if paragraph.paragraph_format.line_spacing is None or paragraph.paragraph_format.line_spacing == 1.0:
                if paragraph.style.name == 'Normal':
                    paragraph.paragraph_format.line_spacing = 2.0
            
            # Alineación justificada para texto normal
            if paragraph.style.name == 'Normal' and paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Sangría solo para párrafos largos sin otra configuración
            if paragraph.style.name == 'Normal' and len(paragraph.text) > 50:
                if paragraph.paragraph_format.first_line_indent is None or paragraph.paragraph_format.first_line_indent == 0:
                    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    
    def crear_tabla_apa(self, doc, filas, columnas):
        """Crea una tabla con formato APA"""
        tabla = doc.add_table(rows=filas, cols=columnas)
        tabla.style = 'Light Grid'
        
        for row in tabla.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
        
        return tabla
    
    def generar_word(self):
        if not self.sujetos:
            messagebox.showwarning("Advertencia", "Debe agregar al menos un sujeto de investigación")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documento Word", "*.docx"), ("Todos los archivos", "*.*")],
            initialfile=f"Debida_Diligencia_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        
        if not filename:
            return
        
        try:
            doc = Document()
            
            # Configurar márgenes primero
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Título principal con formato APA (usando configuración personalizable)
            titulo_texto = self.configuraciones.get('titulo_documento', 'INFORME DE DEBIDA DILIGENCIA')
            titulo = doc.add_heading(titulo_texto, level=0)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            titulo_run = titulo.runs[0]
            titulo_run.font.name = 'Times New Roman'
            titulo_run.font.size = Pt(16)
            titulo_run.font.bold = True
            titulo.paragraph_format.space_after = Pt(24)
            titulo.paragraph_format.space_before = Pt(0)
            
            # Información general
            datos_generales = [
                ('Fecha de Solicitud: ', self.fecha_solicitud.get()),
                ('Usuario Requirente: ', self.usuario_requirente.get()),
                ('Tipo de Solicitud: ', self.tipo_solicitud.get()),
                ('Descripción: ', self.descripcion_solicitud.get())
            ]
            
            for etiqueta, valor in datos_generales:
                p = doc.add_paragraph()
                run1 = p.add_run(etiqueta)
                run1.font.bold = True
                run1.font.name = 'Times New Roman'
                run1.font.size = Pt(12)
                run2 = p.add_run(valor)
                run2.font.name = 'Times New Roman'
                run2.font.size = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.0
            
            # Serapio
            p = doc.add_paragraph()
            run1 = p.add_run('Serapio: ')
            run1.font.bold = True
            run1.font.name = 'Times New Roman'
            run1.font.size = Pt(12)
            run2 = p.add_run(self.serapio.get())
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            
            # Encabezado: Sujeto de Investigación
            encabezado = doc.add_heading('Sujeto de Investigación', level=1)
            enc_run = encabezado.runs[0]
            enc_run.font.name = 'Times New Roman'
            enc_run.font.size = Pt(14)
            enc_run.font.bold = True
            encabezado.paragraph_format.space_before = Pt(12)
            encabezado.paragraph_format.space_after = Pt(12)
            
            # Tabla de sujetos (con columna de ID Gestión)
            tabla = self.crear_tabla_apa(doc, len(self.sujetos) + 1, 4)
            hdr_cells = tabla.rows[0].cells
            hdr_cells[0].text = 'Nombre'
            hdr_cells[1].text = 'Identificación'
            hdr_cells[2].text = 'Descripción'
            hdr_cells[3].text = 'ID Gestión'
            
            for cell in tabla.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            for i, sujeto in enumerate(self.sujetos, 1):
                row_cells = tabla.rows[i].cells
                row_cells[0].text = sujeto['nombre'].title()  # Formato de nombres propios
                row_cells[1].text = sujeto['identificacion']
                row_cells[2].text = sujeto['descripcion']
                row_cells[3].text = sujeto.get('id_gestion', 'N/A')
            
            # Espacio después de la tabla
            p_espacio = doc.add_paragraph()
            p_espacio.paragraph_format.space_after = Pt(12)
            
            # Fuente de información
            p = doc.add_paragraph()
            run = p.add_run('Fuente de Información: ')
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run2 = p.add_run(self.fuente_info.get(1.0, tk.END).strip())
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Sección 1: Objetivo del Informe
            encabezado = doc.add_heading('1. Objetivo del Informe', level=1)
            for run in encabezado.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
            encabezado.paragraph_format.space_before = Pt(12)
            encabezado.paragraph_format.space_after = Pt(12)
            
            p = doc.add_paragraph(self.plantilla_modificable['objetivo'])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 2.0  # Doble espacio según APA
            
            # Sección 2: Resultados de la Investigación
            encabezado = doc.add_heading('2. Resultados de la Investigación', level=1)
            for run in encabezado.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
            encabezado.paragraph_format.space_before = Pt(12)
            encabezado.paragraph_format.space_after = Pt(12)
            
            tabla_resultados = self.crear_tabla_apa(doc, len(self.resultados) + 1, 3)
            hdr_cells = tabla_resultados.rows[0].cells
            hdr_cells[0].text = 'Fuente'
            hdr_cells[1].text = 'Resultado'
            hdr_cells[2].text = 'Fecha'
            
            # Formato de encabezados
            for cell in tabla_resultados.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
            
            # Datos de resultados
            for i, (fuente, widget) in enumerate(self.resultados.items(), 1):
                row_cells = tabla_resultados.rows[i].cells
                row_cells[0].text = fuente
                row_cells[1].text = widget.get()
                row_cells[2].text = 'N/A'
                
                # Formato de celdas
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
            
            # Espacio después de la tabla
            doc.add_paragraph()
            
            # Sección 3: Conclusión
            encabezado = doc.add_heading('3. Conclusión', level=1)
            enc_run = encabezado.runs[0]
            enc_run.font.name = 'Times New Roman'
            enc_run.font.size = Pt(14)
            enc_run.font.bold = True
            encabezado.paragraph_format.space_before = Pt(12)
            encabezado.paragraph_format.space_after = Pt(12)
            
            # Preparar texto de nombres
            if len(self.sujetos) == 1:
                nombres_texto = f"la persona {self.sujetos[0]['nombre'].title()}"
            else:
                nombres = [s['nombre'].title() for s in self.sujetos]  # Formato de nombres propios
                nombres_texto = "las personas " + ", ".join(nombres[:-1]) + f" y {nombres[-1]}"
            
            conclusion = self.plantilla_modificable['conclusion_template'].format(
                nombres=nombres_texto,
                nivel_riesgo=self.nivel_riesgo.get()
            )
            
            # Inciso a
            p = doc.add_paragraph()
            run1 = p.add_run('a. ')
            run1.font.name = 'Times New Roman'
            run1.font.size = Pt(12)
            run1.font.bold = True
            run2 = p.add_run(conclusion)
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.first_line_indent = Inches(0.5)
            
            # Inciso b
            p = doc.add_paragraph()
            run1 = p.add_run('b. ')
            run1.font.name = 'Times New Roman'
            run1.font.size = Pt(12)
            run1.font.bold = True
            run2 = p.add_run(self.plantilla_modificable['compromiso'])
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.first_line_indent = Inches(0.5)
            
            # Incisos adicionales personalizados
            if 'incisos_adicionales' in self.plantilla_modificable and self.plantilla_modificable['incisos_adicionales']:
                for i, texto_inciso in enumerate(self.plantilla_modificable['incisos_adicionales'], start=3):
                    letra = chr(96 + i)  # c, d, e, etc.
                    p = doc.add_paragraph()
                    run1 = p.add_run(f'{letra}. ')
                    run1.font.name = 'Times New Roman'
                    run1.font.size = Pt(12)
                    run1.font.bold = True
                    run2 = p.add_run(texto_inciso)
                    run2.font.name = 'Times New Roman'
                    run2.font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(12)
                    p.paragraph_format.line_spacing = 2.0
                    p.paragraph_format.first_line_indent = Inches(0.5)
            
            # Sección 4: Firma y Lugar de Emisión
            encabezado = doc.add_heading('4. Firma y Lugar de Emisión', level=1)
            enc_run = encabezado.runs[0]
            enc_run.font.name = 'Times New Roman'
            enc_run.font.size = Pt(14)
            enc_run.font.bold = True
            encabezado.paragraph_format.space_before = Pt(12)
            encabezado.paragraph_format.space_after = Pt(12)
            
            meses = {
                1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
                5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
                9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
            }
            dias = {
                0: 'lunes', 1: 'martes', 2: 'miércoles', 3: 'jueves',
                4: 'viernes', 5: 'sábado', 6: 'domingo'
            }
            
            ahora = datetime.now()
            dia_semana = dias[ahora.weekday()]
            dia = ahora.day
            mes = meses[ahora.month]
            año = ahora.year
            fecha_actual = f"{dia_semana} {dia} de {mes} de {año}"
            
            p = doc.add_paragraph(f"Emitido en {self.plantilla_modificable['lugar_emision']}, el {fecha_actual}")
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(24)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Espacio adicional para las firmas (más separación)
            doc.add_paragraph()  # Espacio 1
            doc.add_paragraph()  # Espacio 2
            doc.add_paragraph()  # Espacio 3
            
            # Tabla de firmas profesional (sin bordes) con más ancho
            tabla_firmas = doc.add_table(rows=1, cols=2)
            tabla_firmas.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Eliminar todos los bordes de la tabla
            for row in tabla_firmas.rows:
                for cell in row.cells:
                    cell.width = Inches(3.5)  # Aumentado de 2.5 a 3.5 pulgadas
                    tcPr = cell._element.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'none')
                        tcBorders.append(border)
                    tcPr.append(tcBorders)
            
            # Columna izquierda: Jefe de Cumplimiento
            cell_izq = tabla_firmas.rows[0].cells[0]
            
            # Línea de firma
            p_linea_izq = cell_izq.add_paragraph()
            p_linea_izq.add_run('_' * 35)  # Aumentado de 30 a 35
            p_linea_izq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p_linea_izq.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            p_linea_izq.paragraph_format.space_after = Pt(8)  # Aumentado de 6 a 8
            
            # Nombre del jefe
            jefe = self.plantilla_modificable['jefe_cumplimiento'] if self.plantilla_modificable['jefe_cumplimiento'] else "Jefe de Cumplimiento"
            p_nombre_izq = cell_izq.add_paragraph()
            run_nombre = p_nombre_izq.add_run(jefe)
            run_nombre.font.name = 'Times New Roman'
            run_nombre.font.size = Pt(12)
            run_nombre.font.bold = True
            p_nombre_izq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_nombre_izq.paragraph_format.space_after = Pt(4)  # Aumentado de 3 a 4
            
            # Cargo
            p_cargo_izq = cell_izq.add_paragraph()
            run_cargo = p_cargo_izq.add_run('Jefe de Cumplimiento')
            run_cargo.font.name = 'Times New Roman'
            run_cargo.font.size = Pt(11)
            p_cargo_izq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Columna derecha: Analista
            cell_der = tabla_firmas.rows[0].cells[1]
            
            # Línea de firma
            p_linea_der = cell_der.add_paragraph()
            p_linea_der.add_run('_' * 35)  # Aumentado de 30 a 35
            p_linea_der.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p_linea_der.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            p_linea_der.paragraph_format.space_after = Pt(8)  # Aumentado de 6 a 8
            
            # Nombre del analista
            analista = self.plantilla_modificable['analista'] if self.plantilla_modificable['analista'] else "Analista"
            p_nombre_der = cell_der.add_paragraph()
            run_nombre_der = p_nombre_der.add_run(analista)
            run_nombre_der.font.name = 'Times New Roman'
            run_nombre_der.font.size = Pt(12)
            run_nombre_der.font.bold = True
            p_nombre_der.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_nombre_der.paragraph_format.space_after = Pt(4)  # Aumentado de 3 a 4
            
            # Cargo y firma del responsable designado
            p_cargo_der = cell_der.add_paragraph()
            run_cargo_der = p_cargo_der.add_run('Analista de Cumplimiento')
            run_cargo_der.font.name = 'Times New Roman'
            run_cargo_der.font.size = Pt(11)
            p_cargo_der.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cargo_der.paragraph_format.space_after = Pt(4)  # Aumentado de 3 a 4
            
            p_firma_resp = cell_der.add_paragraph()
            run_firma = p_firma_resp.add_run('Firma del Responsable Designado')
            run_firma.font.name = 'Times New Roman'
            run_firma.font.size = Pt(10)
            run_firma.font.italic = True
            p_firma_resp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Aplicar formato APA al final
            self.aplicar_formato_apa(doc)
            
            # Guardar documento
            doc.save(filename)
            
            self.status_label.config(text=f"✓ Documento Word generado: {os.path.basename(filename)}")
            messagebox.showinfo("Éxito", f"Documento generado exitosamente:\n{filename}")
            
        except PermissionError:
            messagebox.showerror("Error de Permisos", 
                               "No se puede guardar el archivo. Verifique que:\n" +
                               "1. El archivo no esté abierto en Word\n" +
                               "2. Tenga permisos de escritura en la carpeta")
        except Exception as e:
            import traceback
            error_detallado = traceback.format_exc()
            messagebox.showerror("Error", f"Error al generar documento:\n{str(e)}\n\nDetalles:\n{error_detallado[:200]}")
    
    def generar_excel_completo(self):
        if not self.sujetos:
            messagebox.showwarning("Advertencia", "Debe agregar al menos un sujeto de investigación")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Archivo Excel", "*.xlsx"), ("Todos los archivos", "*.*")],
            initialfile=f"Debida_Diligencia_Completo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        )
        
        if not filename:
            return
        
        try:
            meses = {
                1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
                5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
                9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
            }
            dias = {
                0: 'lunes', 1: 'martes', 2: 'miércoles', 3: 'jueves',
                4: 'viernes', 5: 'sábado', 6: 'domingo'
            }
            
            ahora = datetime.now()
            dia_semana = dias[ahora.weekday()]
            dia = ahora.day
            mes = meses[ahora.month]
            año = ahora.year
            fecha_emision = f"{dia_semana} {dia} de {mes} de {año}"
            
            data_dict = {
                'ID Gestión': [],
                'Fecha de Solicitud': [],
                'Usuario Requirente': [],
                'Tipo de Solicitud': [],
                'Descripción': [],
                'Serapio': [],
                'Nombre': [],
                'Identificación': [],
                'Descripción del Sujeto': []
            }
            
            for fuente in self.configuraciones['fuentes_investigacion']:
                data_dict[fuente] = []
            
            data_dict['Nivel de Riesgo'] = []
            data_dict['Fecha de Emisión'] = []
            
            for sujeto in self.sujetos:
                data_dict['ID Gestión'].append(sujeto.get('id_gestion', 'N/A'))
                data_dict['Fecha de Solicitud'].append(self.fecha_solicitud.get())
                data_dict['Usuario Requirente'].append(self.usuario_requirente.get())
                data_dict['Tipo de Solicitud'].append(self.tipo_solicitud.get())
                data_dict['Descripción'].append(self.descripcion_solicitud.get())
                data_dict['Serapio'].append(self.serapio.get())
                data_dict['Nombre'].append(sujeto['nombre'].title())  # Formato nombres propios
                data_dict['Identificación'].append(sujeto['identificacion'])
                data_dict['Descripción del Sujeto'].append(sujeto['descripcion'])
                
                for fuente in self.configuraciones['fuentes_investigacion']:
                    if fuente in self.resultados:
                        data_dict[fuente].append(self.resultados[fuente].get())
                    else:
                        data_dict[fuente].append('N/A')
                
                data_dict['Nivel de Riesgo'].append(self.nivel_riesgo.get().upper())
                data_dict['Fecha de Emisión'].append(fecha_emision)
            
            df = pd.DataFrame(data_dict)
            
            with pd.ExcelWriter(filename, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Debidas Diligencias', index=False)
                
                workbook = writer.book
                worksheet = writer.sheets['Debidas Diligencias']
                
                from openpyxl.styles import Font, PatternFill, Alignment, Border
                
                header_font = Font(name='Calibri', size=12, bold=True, color='FFFFFF')
                header_fill = PatternFill(start_color='2C3E50', end_color='2C3E50', fill_type='solid')
                header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                
                data_font = Font(name='Calibri', size=11)
                data_alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                
                fill_even = PatternFill(start_color='FFFFFF', end_color='FFFFFF', fill_type='solid')
                fill_odd = PatternFill(start_color='F8F9FA', end_color='F8F9FA', fill_type='solid')
                
                no_border = Border()
                
                for cell in worksheet[1]:
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = header_alignment
                    cell.border = no_border
                
                for idx, row in enumerate(worksheet.iter_rows(min_row=2, max_row=worksheet.max_row, min_col=1, max_col=worksheet.max_column), start=2):
                    if idx % 2 == 0:
                        row_fill = fill_even
                    else:
                        row_fill = fill_odd
                    
                    for cell in row:
                        cell.font = data_font
                        cell.alignment = data_alignment
                        cell.border = no_border
                        cell.fill = row_fill
                
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    
                    for cell in column:
                        try:
                            if cell.value:
                                cell_length = len(str(cell.value))
                                if cell_length > max_length:
                                    max_length = cell_length
                        except:
                            pass
                    
                    adjusted_width = min(max_length + 2, 50)
                    adjusted_width = max(adjusted_width, 12)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
                
                worksheet.freeze_panes = 'A2'
                worksheet.auto_filter.ref = worksheet.dimensions
                worksheet.row_dimensions[1].height = 30
            
            self.status_label.config(text=f"✓ Excel completo generado: {os.path.basename(filename)}")
            messagebox.showinfo("Éxito", f"Excel generado exitosamente:\n{filename}\n\nCada fila representa una debida diligencia completa.")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error al generar Excel: {str(e)}")
    
    def generar_excel_pestanas(self):
        if not self.sujetos:
            messagebox.showwarning("Advertencia", "Debe agregar al menos un sujeto de investigación")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Archivo Excel", "*.xlsx"), ("Todos los archivos", "*.*")],
            initialfile=f"Debida_Diligencia_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        )
        
        if not filename:
            return
        
        try:
            with pd.ExcelWriter(filename, engine='openpyxl') as writer:
                resumen_data = {
                    'Campo': ['ID Gestión', 'Fecha de Solicitud', 'Usuario Requirente', 'Tipo de Solicitud', 'Descripción', 'Serapio'],
                    'Valor': [
                        self.id_gestion.get(),
                        self.fecha_solicitud.get(),
                        self.usuario_requirente.get(),
                        self.tipo_solicitud.get(),
                        self.descripcion_solicitud.get(),
                        self.serapio.get()
                    ]
                }
                df_resumen = pd.DataFrame(resumen_data)
                df_resumen.to_excel(writer, sheet_name='Resumen', index=False)
                
                for i, sujeto in enumerate(self.sujetos, 1):
                    data = {
                        'Campo': ['Nombre', 'Identificación', 'Descripción', 'ID Gestión', '', 'Resultados:'],
                        'Valor': [sujeto['nombre'].title(), sujeto['identificacion'], sujeto['descripcion'],  # Formato nombres propios
                                 sujeto.get('id_gestion', 'N/A'), '', '']
                    }
                    
                    for fuente, widget in self.resultados.items():
                        data['Campo'].append(fuente)
                        data['Valor'].append(widget.get())
                    
                    data['Campo'].append('Nivel de Riesgo')
                    data['Valor'].append(self.nivel_riesgo.get())
                    
                    df = pd.DataFrame(data)
                    sheet_name = sujeto['nombre'][:31]
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            self.status_label.config(text=f"✓ Excel generado: {os.path.basename(filename)}")
            messagebox.showinfo("Éxito", f"Excel generado exitosamente:\n{filename}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error al generar Excel: {str(e)}")
    
    def generar_todo(self):
        self.generar_word()
        self.generar_excel_completo()
        self.generar_excel_pestanas()

if __name__ == "__main__":
    root = tk.Tk()
    app = DueDiligenceSystem(root)
    root.mainloop()